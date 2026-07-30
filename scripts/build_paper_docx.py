"""日本語研究草稿Markdownを査読用DOCXへ変換する。"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

CONTENT_WIDTH_DXA = 9_360
TABLE_INDENT_DXA = 120
BODY_FONT = "Calibri"
CJK_FONT = "Yu Gothic"
MATH_FONT = "Cambria Math"
HEADING_BLUE = "2E74B5"
HEADING_DARK_BLUE = "1F4D78"
INK_BLUE = "0B2545"
MUTED = "5B6573"
TABLE_FILL = "F4F6F9"
TABLE_BORDER = "CBD2DA"


def build_document(source_path: Path, output_path: Path) -> None:
    """narrative_proposal presetを適用して論文草稿を生成する。"""

    document = Document()
    _configure_page(document)
    _configure_styles(document)
    bullet_number_id = _create_bullet_numbering(document)
    _set_header_footer(document)
    _render_markdown(
        document=document,
        source_text=source_path.read_text(encoding="utf-8"),
        bullet_number_id=bullet_number_id,
    )
    _set_document_properties(document)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    _set_style_font(normal, BODY_FONT, CJK_FONT, 11, "000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _configure_heading(
        document.styles["Heading 1"],
        size=16,
        color=HEADING_BLUE,
        before=18,
        after=10,
    )
    _configure_heading(
        document.styles["Heading 2"],
        size=13,
        color=HEADING_BLUE,
        before=12,
        after=6,
    )
    _configure_heading(
        document.styles["Heading 3"],
        size=12,
        color=HEADING_DARK_BLUE,
        before=8,
        after=4,
    )

    title = document.styles.add_style("Academic Title", 1)
    _set_style_font(title, BODY_FONT, CJK_FONT, 24, INK_BLUE, bold=True)
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.keep_with_next = True

    subtitle = document.styles.add_style("Academic Subtitle", 1)
    _set_style_font(
        subtitle,
        BODY_FONT,
        CJK_FONT,
        14,
        HEADING_DARK_BLUE,
    )
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.keep_with_next = True

    metadata = document.styles.add_style("Draft Metadata", 1)
    _set_style_font(metadata, BODY_FONT, CJK_FONT, 9.5, MUTED)
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(2)
    metadata.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    keywords = document.styles.add_style("Keywords", 1)
    _set_style_font(keywords, BODY_FONT, CJK_FONT, 9.5, MUTED)
    keywords.paragraph_format.space_before = Pt(4)
    keywords.paragraph_format.space_after = Pt(10)
    keywords.paragraph_format.line_spacing = 1.167

    equation = document.styles.add_style("Display Equation", 1)
    _set_style_font(equation, MATH_FONT, CJK_FONT, 10.5, "000000")
    equation.paragraph_format.space_before = Pt(6)
    equation.paragraph_format.space_after = Pt(8)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.keep_together = True

    bibliography = document.styles.add_style("Bibliography", 1)
    _set_style_font(bibliography, BODY_FONT, CJK_FONT, 9.5, "222222")
    bibliography.paragraph_format.left_indent = Inches(0.3)
    bibliography.paragraph_format.first_line_indent = Inches(-0.3)
    bibliography.paragraph_format.space_before = Pt(0)
    bibliography.paragraph_format.space_after = Pt(5)
    bibliography.paragraph_format.line_spacing = 1.167

    table_text = document.styles.add_style("Scientific Table Text", 1)
    _set_style_font(table_text, BODY_FONT, CJK_FONT, 9, "111111")
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(2)
    table_text.paragraph_format.line_spacing = 1.0


def _configure_heading(
    style,
    *,
    size: float,
    color: str,
    before: float,
    after: float,
) -> None:
    _set_style_font(style, BODY_FONT, CJK_FONT, size, color, bold=True)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.keep_together = True


def _set_style_font(
    style,
    latin_font: str,
    east_asia_font: str,
    size: float,
    color: str,
    *,
    bold: bool = False,
) -> None:
    style.font.name = latin_font
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin_font)
    fonts.set(qn("w:hAnsi"), latin_font)
    fonts.set(qn("w:eastAsia"), east_asia_font)


def _set_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5),
        WD_TAB_ALIGNMENT.RIGHT,
    )
    left_run = paragraph.add_run("Reservoir Dynamics Research")
    _format_run(left_run, 8.5, MUTED)
    paragraph.add_run("\t")
    right_run = paragraph.add_run("研究草稿")
    _format_run(right_run, 8.5, MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)
    _add_page_number(footer_paragraph)


def _add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    _format_run(run, 8.5, MUTED)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")
    field_text = OxmlElement("w:t")
    field_text.text = "1"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend(
        [field_begin, instruction, field_separate, field_text, field_end]
    )


def _render_markdown(
    *,
    document: Document,
    source_text: str,
    bullet_number_id: int,
) -> None:
    lines = source_text.splitlines()
    line_index = 0
    title_rendered = False
    subtitle_rendered = False
    while line_index < len(lines):
        line = lines[line_index].strip()
        if not line:
            line_index += 1
            continue
        if line.startswith("|") and line_index + 1 < len(lines):
            table_lines: list[str] = []
            while line_index < len(lines) and lines[line_index].strip().startswith(
                "|"
            ):
                table_lines.append(lines[line_index].strip())
                line_index += 1
            _add_markdown_table(document, table_lines)
            continue
        if line == r"\[":
            equation_lines: list[str] = []
            line_index += 1
            while line_index < len(lines) and lines[line_index].strip() != r"\]":
                equation_lines.append(lines[line_index].strip())
                line_index += 1
            line_index += 1
            _add_equation(document, equation_lines)
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Academic Title")
            _add_inline_runs(paragraph, line[2:].strip())
            title_rendered = True
            line_index += 1
            continue
        if line.startswith("## "):
            heading_text = line[3:].strip()
            if title_rendered and not subtitle_rendered:
                paragraph = document.add_paragraph(
                    style="Academic Subtitle"
                )
                _add_inline_runs(paragraph, heading_text)
                subtitle_rendered = True
            else:
                document.add_heading(heading_text, level=1)
            line_index += 1
            continue
        if line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
            line_index += 1
            continue
        if line.startswith("- "):
            while line_index < len(lines) and lines[line_index].strip().startswith(
                "- "
            ):
                paragraph = document.add_paragraph()
                _apply_bullet(paragraph, bullet_number_id)
                _add_inline_runs(
                    paragraph,
                    lines[line_index].strip()[2:].strip(),
                )
                line_index += 1
            continue
        if re.match(r"^\[\d+\]\s", line):
            paragraph = document.add_paragraph(style="Bibliography")
            _add_inline_runs(paragraph, line)
            line_index += 1
            continue
        paragraph_lines = [line]
        line_index += 1
        while line_index < len(lines):
            next_line = lines[line_index].strip()
            if (
                not next_line
                or next_line.startswith("#")
                or next_line.startswith("|")
                or next_line.startswith("- ")
                or next_line == r"\["
                or re.match(r"^\[\d+\]\s", next_line)
            ):
                break
            paragraph_lines.append(next_line)
            line_index += 1
        text = " ".join(paragraph_lines).replace("  ", " ")
        style_name = (
            "Draft Metadata"
            if text.startswith("研究草稿") or re.match(r"^20\d{2}年", text)
            else "Keywords"
            if text.startswith("キーワード:")
            else None
        )
        paragraph = document.add_paragraph(style=style_name)
        _add_inline_runs(paragraph, text)


def _add_inline_runs(paragraph, text: str) -> None:
    token_pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\\\(.*?\\\))")
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            _add_text_run(paragraph, text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = _add_text_run(paragraph, token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = _add_text_run(paragraph, token[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
            run.font.size = Pt(9.5)
        else:
            run = _add_text_run(
                paragraph,
                _plain_math(token[2:-2]),
            )
            run.font.name = MATH_FONT
            run._element.rPr.rFonts.set(qn("w:ascii"), MATH_FONT)
            run._element.rPr.rFonts.set(qn("w:hAnsi"), MATH_FONT)
        position = match.end()
    if position < len(text):
        _add_text_run(paragraph, text[position:])


def _add_text_run(paragraph, text: str):
    run = paragraph.add_run(text)
    _format_run(run, None, None)
    return run


def _format_run(
    run,
    size: float | None,
    color: str | None,
) -> None:
    run.font.name = BODY_FONT
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), BODY_FONT)
    fonts.set(qn("w:hAnsi"), BODY_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_equation(document: Document, equation_lines: list[str]) -> None:
    text = "\n".join(
        _plain_math(line)
        for line in equation_lines
        if line not in (r"\begin{aligned}", r"\end{aligned}")
    )
    paragraph = document.add_paragraph(style="Display Equation")
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        run = paragraph.add_run(line)
        run.font.name = MATH_FONT
        run._element.rPr.rFonts.set(qn("w:ascii"), MATH_FONT)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), MATH_FONT)
        run.font.size = Pt(10.5)


def _plain_math(text: str) -> str:
    replacements = (
        (r"\left", ""),
        (r"\right", ""),
        (r"\boldsymbol{x}", "x"),
        (r"\boldsymbol{y}", "y"),
        (r"\boldsymbol{s}", "s"),
        (r"\boldsymbol{d}", "d"),
        (r"\boldsymbol{\eta}", "η"),
        (r"\mathcal{A}", "A"),
        (r"\mathcal{B}", "B"),
        (r"\mathbf{1}", "1"),
        (r"\operatorname{atanh}", "atanh"),
        (r"\bar{\mu}", "μ̄"),
        (r"\bar\mu", "μ̄"),
        (r"\tanh", "tanh"),
        (r"\infty", "∞"),
        (r"\qquad", "  "),
        (r"\quad", " "),
        (r"\times", "×"),
        (r"\sum", "Σ"),
        (r"\max", "max"),
        (r"\min", "min"),
        (r"\rho", "ρ"),
        (r"\mu", "μ"),
        (r"\eta", "η"),
        (r"\ge", "≥"),
        (r"\le", "≤"),
        (r"\in", "∈"),
        (r"\|", "‖"),
        (r"\{", "{"),
        (r"\}", "}"),
        (r"\\", ""),
        ("&", ""),
    )
    output = text
    for source, target in replacements:
        output = output.replace(source, target)
    output = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", output)
    output = re.sub(r"\\text\{([^{}]+)\}", r"\1", output)
    output = re.sub(r"\\mathcal\{([^{}]+)\}", r"\1", output)
    output = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", output)
    output = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", output)
    output = re.sub(r"\\([A-Za-z]+)", r"\1", output)
    output = output.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", output).strip()


def _add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [
        tuple(cell.strip() for cell in line.strip("|").split("|"))
        for line in lines
    ]
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[2:]
    column_count = len(header)
    widths = _table_widths(column_count)
    table = document.add_table(rows=1, cols=column_count)
    table.autofit = False
    _set_table_geometry(table, widths)
    _write_table_row(table.rows[0], header, widths, is_header=True)
    _mark_repeat_header(table.rows[0])
    for values in body:
        row = table.add_row()
        _write_table_row(row, values, widths, is_header=False)
    spacing = document.add_paragraph()
    spacing.paragraph_format.space_before = Pt(0)
    spacing.paragraph_format.space_after = Pt(2)


def _table_widths(column_count: int) -> tuple[int, ...]:
    patterns = {
        3: (4_000, 1_800, 3_560),
        4: (2_350, 2_850, 2_100, 2_060),
        5: (2_800, 1_640, 1_640, 1_640, 1_640),
    }
    if column_count in patterns:
        return patterns[column_count]
    base_width = CONTENT_WIDTH_DXA // column_count
    widths = [base_width] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return tuple(widths)


def _set_table_geometry(table, widths: tuple[int, ...]) -> None:
    table_properties = table._tbl.tblPr
    _set_child_value(table_properties, "w:tblW", "w:w", CONTENT_WIDTH_DXA)
    table_properties.find(qn("w:tblW")).set(qn("w:type"), "dxa")
    _set_child_value(
        table_properties,
        "w:tblInd",
        "w:w",
        TABLE_INDENT_DXA,
    )
    table_properties.find(qn("w:tblInd")).set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_properties.append(layout)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), TABLE_BORDER)
        borders.append(border)
    table_properties.append(borders)
    margins = OxmlElement("w:tblCellMar")
    for edge, value in (
        ("top", 80),
        ("bottom", 80),
        ("start", 120),
        ("end", 120),
    ):
        margin = OxmlElement(f"w:{edge}")
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    table_properties.append(margins)
    grid = table._tbl.tblGrid
    for child in tuple(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)


def _set_child_value(parent, tag: str, attribute: str, value: int) -> None:
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    child.set(qn(attribute), str(value))


def _write_table_row(
    row,
    values: tuple[str, ...],
    widths: tuple[int, ...],
    *,
    is_header: bool,
) -> None:
    for column_index, (cell, value, width) in enumerate(
        zip(row.cells, values, widths, strict=True)
    ):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
        cell_width.set(qn("w:w"), str(width))
        cell_width.set(qn("w:type"), "dxa")
        if is_header:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), TABLE_FILL)
            cell._tc.get_or_add_tcPr().append(shading)
        paragraph = cell.paragraphs[0]
        paragraph.style = "Scientific Table Text"
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
            if column_index == 0
            else WD_ALIGN_PARAGRAPH.CENTER
        )
        _add_inline_runs(paragraph, value)
        if is_header:
            for run in paragraph.runs:
                run.bold = True


def _mark_repeat_header(row) -> None:
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    row._tr.get_or_add_trPr().append(table_header)


def _create_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    number_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    number_id = max(number_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "260")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "280")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_properties.extend([tabs, indent, spacing])
    level.extend(
        [
            start,
            number_format,
            level_text,
            level_justification,
            paragraph_properties,
        ]
    )
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(number_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return number_id


def _apply_bullet(paragraph, number_id: int) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    paragraph_properties = paragraph._p.get_or_add_pPr()
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(number_id))
    number_properties.extend([level, number])
    paragraph_properties.append(number_properties)


def _set_document_properties(document: Document) -> None:
    properties = document.core_properties
    properties.title = "アトラクタ数を越えたリザバー評価"
    properties.subject = "ロバスト・レパートリー余裕と外乱下記憶性能"
    properties.keywords = (
        "reservoir computing; attractor; robust invariant set; RNN"
    )
    properties.comments = "研究草稿 v0.1"


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    arguments = parser.parse_args()
    build_document(arguments.source, arguments.output)


if __name__ == "__main__":
    main()
