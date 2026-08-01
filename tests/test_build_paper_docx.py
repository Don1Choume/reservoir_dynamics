from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from scripts.build_paper_docx import (
    _add_equation,
    _configure_page,
    _configure_styles,
    _plain_math,
    _set_header_footer,
    build_document,
)


def test_common_math_indices_are_readable_unicode() -> None:
    rendered = _plain_math(
        r"\boldsymbol{x}_{t+1}=\tanh("
        r"W\boldsymbol{x}_t+\boldsymbol{\eta}_t),"
        r"\qquad \|\boldsymbol{\eta}_t\|_\infty\le e."
    )

    assert rendered == "xₜ₊₁=tanh(Wxₜ+ηₜ), ‖ηₜ‖∞≤ e."


def test_integral_is_converted_before_set_membership() -> None:
    rendered = _plain_math(
        r"\bar\mu/e=\int_0^\infty R(eu)\,du"
    )

    assert rendered == "μ̄/e=∫₀^∞ R(eu) du"


def test_multiline_display_equation_is_compacted_to_one_line() -> None:
    document = Document()
    _configure_styles(document)

    _add_equation(
        document,
        [
            r"\begin{aligned}",
            r"\mu_{\boldsymbol{s}}",
            "=",
            r"\max_{0<m<1}",
            r"\min_i",
            "[",
            r"P_i m+N_i-\operatorname{atanh}(m)",
            "]",
            r"\end{aligned}",
        ],
    )

    equation = document.paragraphs[-1]
    assert "\n" not in equation.text
    assert "μₛ" in equation.text
    assert "minᵢ" in equation.text
    assert "Pᵢ" in equation.text


def test_document_styles_reserve_footer_space_and_left_align_sources() -> None:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _set_header_footer(document)

    assert document.sections[0].bottom_margin.inches == pytest.approx(1.2)
    assert document.settings.odd_and_even_pages_header_footer is True
    assert (
        document.sections[0].even_page_header.paragraphs[0].text
        == "Reservoir Dynamics Research\t研究草稿"
    )
    assert (
        document.styles["Bibliography"].paragraph_format.alignment
        == WD_ALIGN_PARAGRAPH.LEFT
    )
    assert (
        document.styles["Bibliography"].paragraph_format.space_after.pt
        == pytest.approx(3)
    )
    assert (
        document.styles["Reproducibility"].paragraph_format.alignment
        == WD_ALIGN_PARAGRAPH.LEFT
    )


def test_reproducibility_paragraph_uses_nonjustified_style(
    tmp_path: Path,
) -> None:
    source_path = tmp_path / "paper.md"
    output_path = tmp_path / "paper.docx"
    source_path.write_text(
        "\n".join(
            (
                "# 題名",
                "## 副題",
                "研究草稿 v0.1",
                "## データ・コードと再現性",
                (
                    "実験spec、seed、判定、導出済みartifactを保存した。"
                    "source manifest SHA-256 `abcdef` で固定した。"
                ),
                "## 参考文献",
                "[1] Example. https://example.test",
            )
        ),
        encoding="utf-8",
    )

    build_document(source_path, output_path)

    rendered = Document(output_path)
    reproducibility = next(
        paragraph
        for paragraph in rendered.paragraphs
        if paragraph.text.startswith("実験spec")
    )
    bibliography = next(
        paragraph
        for paragraph in rendered.paragraphs
        if paragraph.text.startswith("[1]")
    )
    assert reproducibility.style.name == "Reproducibility"
    assert reproducibility.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert bibliography.alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_pagebreak_marker_creates_word_page_break(tmp_path: Path) -> None:
    source_path = tmp_path / "paper.md"
    output_path = tmp_path / "paper.docx"
    source_path.write_text(
        "\n".join(
            (
                "# 題名",
                "## 副題",
                "研究草稿 v0.2",
                "本文前半。",
                "",
                "<!-- pagebreak -->",
                "",
                "### 後半",
                "本文後半。",
            )
        ),
        encoding="utf-8",
    )

    build_document(source_path, output_path)

    rendered = Document(output_path)
    assert "<!-- pagebreak -->" not in "\n".join(
        paragraph.text for paragraph in rendered.paragraphs
    )
    assert 'w:type="page"' in rendered._element.xml
